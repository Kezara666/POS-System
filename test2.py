from tkinter import *
from tkinter import ttk
import tkinter as tk
import sqlite3
from docx2pdf import convert
from tkinter import messagebox
from docx import Document
import time
from datetime import date
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx
import os
import shutil
import docx
from docx.shared import Mm
import profit as profite
Total_QWpr = 0.0
baseekata_danna_Q_WPR = 0.0
MULUWIKUNUM = 0.0
shudda_labaya = 0.0



root = Tk()
root.title("NETHMAL SUPERT MARKET BILLING SYSTEM")

root.iconbitmap("img/a.ico")
root.config(background='green')
root.geometry('720x720')
root.resizable(width=False, height=False)
bg = PhotoImage(file="img/cover.png")
l = Label(root, image=bg)
l.place(x=0, y=0)


#######################################################################################add karana eke

def entru():
    from tkinter import messagebox

    ads = Toplevel()
    ads.title("METHMAL SUPERT MARKET BILLING SYSTEM")

    ads.iconbitmap("img/a.ico")
    ads.config(background='green')
    ads.geometry('1024x768')
    ads.resizable(width=False, height=False)
    bg = PhotoImage(file="img/3cover.png")
    l = Label(ads, image=bg)
    l.place(x=0, y=0)
    ############BUTTOUN LABEL

    code = Entry(ads, bd="0", font="Helvetica 18 bold", width="13", fg="black", justify="center", bg='#697c95')
    code.place(x=185, y=77)
    #####
    NAME = Entry(ads, bd="0", font="Helvetica 18 bold", width="28", fg="black", justify="center", bg='#697c95')
    NAME.place(x=180, y=139)
    ###########
    #####

    #######

    ###
    #####
    QNT = Entry(ads, bd="0", font="Helvetica 18 bold", width="10", fg="black", justify="center", bg='#697c95')
    QNT.place(x=180, y=205)
    #####

    ###
    WP = Entry(ads, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#697c95')
    WP.place(x=188, y=271)
    #####
    RETAP = Entry(ads, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#697c95')
    RETAP.place(x=190, y=335)
    #####

    #####
    discoun = Entry(ads, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#697c95')
    discoun.place(x=190, y=394)

    def View():
        con1 = sqlite3.connect("item.db")

        cur1 = con1.cursor()

        cur1.execute("SELECT * FROM item2")

        rows = cur1.fetchall()

        for row in rows:
            print(row)

            tree.insert("", tk.END, values=row)

        con1.close()

    def removeall():
        for recode in tree.get_children():
            tree.delete(recode)

    style = ttk.Style()
    style.configure("mystyle.Treeview", background="black",
                    fieldbackground="black", foreground="white")
    style.configure("mystyle.Treeview", highlightthickness=0, bd=0,
                    font=('Calibri', 11)),  # Modify the font of the body
    style.configure("mystyle.Treeview.Heading", font=('Calibri', 13, 'bold'))  # Modify the font of the headings
    style.layout("mystyle.Treeview", [('mystyle.Treeview.treearea', {'sticky': 'nswe'})])  # Remove the borders

    fram = Frame(ads)
    fram.place(x=2, y=454)
    trs = Scrollbar(fram)
    trs.pack(side=LEFT, fill=Y)


    tree = ttk.Treeview(fram, column=("c1", "c2", "c3", "c4", "c5", "c6"), show='headings', style="mystyle.Treeview",
                        yscrollcommand=trs.set)
    trs.config(command=tree.yview)

    tree.column("#1", anchor=tk.CENTER, width=100)

    tree.heading("#1", text="BARCODE")

    tree.column("#2", anchor=tk.CENTER, width=300)

    tree.heading("#2", text="DISCRIPTION")

    tree.column("#3", anchor=tk.CENTER, width=70)

    tree.heading("#3", text="QNT", )

    tree.column("#4", anchor=tk.CENTER)

    tree.heading("#4", text="Purchase price")

    tree.column("#5", anchor=tk.CENTER)

    tree.heading("#5", text="DISCOUNT")
    tree.column("#5", minwidth=0, width=150, stretch=NO)

    tree.column("#6", anchor=tk.CENTER)

    tree.heading("#6", text="Retail Price")

    tree.pack()
    View()

    def save():
        try:
            conn = sqlite3.connect('item.db')
            sr = code.get()
            print("HRI conn")
            c = conn.cursor()
            print("HRI cor")
            c.execute("DELETE FROM item2 WHERE NUMBER ='%s'" % sr)
            print("HRI cor")
            conn.commit()
            conn.close()

            ar = float(code.get()), str(NAME.get()), float(QNT.get()), float(WP.get()), float(discoun.get()), float(RETAP.get())
            removeall()
            conn = sqlite3.connect('item.db')
            c = conn.cursor()
            c.execute("INSERT INTO item VALUES (:num, :nama, :qnt, :thoga,:sil,:dis)",
                      {

                          'num': float(code.get()),
                          'nama': str(NAME.get()),
                          'qnt': float(QNT.get()),
                          'thoga': float(WP.get()),
                          'sil': float(discoun.get()),
                          'dis': float(RETAP.get())

                      }

                      )
            conn.commit()
            conn.close()
            removeall()
            View()

            code.delete(0, END)
            NAME.delete(0, END)
            WP.delete(0, END)
            RETAP.delete(0, END)
            QNT.delete(0, END)
            discoun.delete(0, END)






        except:
            removeall()
            conn = sqlite3.connect('item.db')
            c = conn.cursor()
            c.execute("INSERT INTO item2 VALUES (:num, :nama, :qnt, :thoga,:sil,:dis)",
                      {

                          'num': float(code.get()),
                          'nama': str(NAME.get()),
                          'qnt': float(QNT.get()),
                          'thoga': float(WP.get()),
                          'sil': float(discoun.get()),
                          'dis': float(RETAP.get())

                      }

                      )
            conn.commit()
            conn.close()
            removeall()
            View()

            code.delete(0, END)
            NAME.delete(0, END)
            WP.delete(0, END)
            RETAP.delete(0, END)
            QNT.delete(0, END)
            discoun.delete(0, END)

    ###########################################################################

    def select(event):
        code.delete(0, END)
        selected = tree.focus()
        val = tree.item(selected, "values")
        code.insert(0, val[0])

        NAME.delete(0, END)
        selected = tree.focus()
        val = tree.item(selected, "values")
        NAME.insert(0, val[1])

        QNT.delete(0, END)
        selected = tree.focus()
        val = tree.item(selected, "values")
        QNT.insert(0, val[2])

        WP.delete(0, END)
        selected = tree.focus()
        val = tree.item(selected, "values")
        WP.insert(0, val[3])
        RETAP.delete(0, END)
        RETAP.insert(0, float(val[5]))

        discoun.delete(0, END)
        discoun.insert(0, float(val[4]))

    def remover():
        x = tree.selection()[0]
        tree.delete(x)
        conn = sqlite3.connect('item.db')
        sr = code.get()
        print("HRI conn")
        c = conn.cursor()
        print("HRI cor")
        c.execute("DELETE FROM item2 WHERE NUMBER ='%s'" % sr)
        print("HRI cor")
        conn.commit()
        conn.close()
        removeall()
        View()
        code.delete(0, END)
        NAME.delete(0, END)
        WP.delete(0, END)
        RETAP.delete(0, END)
        QNT.delete(0, END)
        discoun.delete(0, END)

    delett = PhotoImage(file='img/del.png')
    b2 = Button(ads, image=delett, highlightthickness=0, bd=0, compound=LEFT, command=remover)
    b2.place(x=699, y=269)

    oka = PhotoImage(file="img/save.png")
    b1 = Button(ads, image=oka, highlightthickness=0, bd=0, compound=LEFT, command=save)
    b1.place(x=695, y=166)

    ################tree
    tree.bind("<ButtonRelease-1>", select)
    ######################################################3tree
    ads.mainloop()


def inv():
    list = Toplevel()
    list.title("METHMAL SUPERT MARKET BILLING SYSTEM")

    list.iconbitmap("img/a.ico")
    list.config(background='green')
    list.geometry('1024x768')
    list.resizable(width=False, height=False)
    bg = PhotoImage(file="img/2cover.png")
    l = Label(list, image=bg)
    l.place(x=0, y=0)

    #####
    bar = Entry(list, bd="0", bg='#06a991', font="Helvetica 18 bold", width="25", fg="black", justify="center")
    bar.place(x=168, y=85)
    #######
    dis = Entry(list, bd="0", font="Helvetica 15 bold", width="33", fg="black", justify="center", bg='#06a991')
    dis.place(x=162, y=147)

    #####
    ###

    #####
    normal = Entry(list, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#06a991')
    normal.place(x=720, y=90)
    #####

    ###
    apep = Entry(list, bd="0", fg="black", font="Helvetica 18 bold", width="9", justify="center", bg='#06a991')
    apep.place(x=750, y=150)
    #####
    total = Entry(list, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#06a991')
    total.place(x=750, y=205)
    ########################

    #####
    qnt = Entry(list, bd="0", font="Helvetica 18 bold", width="9", fg="black", justify="center", bg='#06a991')
    qnt.place(x=165, y=200, )

    def calc():

        try:
            conn = sqlite3.connect('inv.db')
            sr = dis.get()
            print("HRI conn")
            c = conn.cursor()
            print("HRI cor")
            c.execute("DELETE FROM inv3 WHERE name ='%s'" % sr)
            print("HRI cor")
            conn.commit()
            conn.close()

            nam = (dis.get())
            q = float(qnt.get())
            norml_price = float(normal.get())
            ape_mila = float(apep.get())
            t = float(total.get())

            qr = float(q * norml_price)
            qd = float(q * ape_mila)

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("INSERT INTO inv3 VALUES ('%s','%s','%s','%s','%s')" % (nam, q, qr, qd, t))
            print('connected')
            conn.commit()
            conn.close()

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("SELECT SUM(qr) FROM inv3 ")
            x = float(c.fetchone()[0])

            c.execute("SELECT SUM(qd) FROM inv3 ")
            y = float(c.fetchone()[0])

            c.execute("SELECT SUM(qt) FROM inv3 ")
            z = float(c.fetchone()[0])

            y = x - y

            print(x, y, z)
            conn.commit()
            conn.close()

            retl = Label(list, bd="0", fg="black", text='    %d' % x, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            retl.place(x=880, y=644)

            discl = Label(list, bd="0", fg="black", text='    %d' % y, font="Helvetica 17 bold", justify="right",
                          bg='#f69679')
            discl.place(x=902, y=684)

            tota = Label(list, bd="0", fg="black", text='    %d' % z, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            tota.place(x=880, y=724)



        except:
            nam = (dis.get())
            q = float(qnt.get())
            norml_price = float(normal.get())
            ape_mila = float(apep.get())
            t = float(total.get())

            qr = float(q * norml_price)
            qd = float(q * ape_mila)

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("INSERT INTO inv3 VALUES ('%s','%s','%s','%s','%s')" % (nam, q, qr, qd, t))
            print('connected')
            conn.commit()
            conn.close()

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("SELECT SUM(qr) FROM inv3 ")
            x = float(c.fetchone()[0])

            c.execute("SELECT SUM(qd) FROM inv3 ")
            y = float(c.fetchone()[0])

            c.execute("SELECT SUM(qt) FROM inv3 ")
            z = float(c.fetchone()[0])

            y = x - y

            print(x, y, z)
            conn.commit()
            conn.close()

            retl = Label(list, bd="0", fg="black", text='    %d' % x, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            retl.place(x=880, y=644)

            discl = Label(list, bd="0", fg="black", text='    %d' % y, font="Helvetica 17 bold", justify="right",
                          bg='#f69679')
            discl.place(x=902, y=684)

            tota = Label(list, bd="0", fg="black", text='    %d' % z, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            tota.place(x=880, y=724)

        conn = sqlite3.connect('inv.db')
        c = conn.cursor()
        print('connected')
        c.execute("SELECT SUM(qr) FROM inv3 ")

        print('qr')
        conn.commit()
        conn.close()

        #

    def View():
        try:
            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("SELECT SUM(qr) FROM inv3 ")
            x = float(c.fetchone()[0])

            c.execute("SELECT SUM(qd) FROM inv3 ")
            y = float(c.fetchone()[0])

            c.execute("SELECT SUM(qt) FROM inv3 ")
            z = float(c.fetchone()[0])

            y = float(x - y)

            print(str(x), y, z)
            conn.commit()
            conn.close()

            retl = Label(list, bd="0", fg="black", text='    %s' % x, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            retl.place(x=880, y=644)

            discl = Label(list, bd="0", fg="black", text='    %s' % y, font="Helvetica 17 bold", justify="right",
                          bg='#f69679')
            discl.place(x=902, y=684)

            tota = Label(list, bd="0", fg="black", text='    %s' % z, font="Helvetica 17 bold", justify="right",
                         bg='#f69679')
            tota.place(x=880, y=724)


        except:
            pass

        con1 = sqlite3.connect("inv.db")

        cur1 = con1.cursor()

        cur1.execute("SELECT * FROM inv2")

        rows = cur1.fetchall()

        for row in rows:
            print(row)

            tree.insert("", tk.END, values=row)

        con1.close()
        ' k'

    style = ttk.Style()
    style.configure("mystyle.Treeview", highlightthickness=0, bd=0,
                    font=('Calibri', 11))  # Modify the font of the body
    style.configure("mystyle.Treeview", background='black')
    style.configure("mystyle.Treeview.Heading", font=('Calibri', 13, 'bold'))  # Modify the font of the headings
    style.layout("mystyle.Treeview", [('mystyle.Treeview.treearea', {'sticky': 'nswe'})])  # Remove the borders

    fram = Frame(list)
    fram.place(x=10, y=285)
    trs = Scrollbar(fram)
    trs.pack(side=LEFT, fill=Y)
    tree = ttk.Treeview(fram, column=("c1", "c2", "c3", "c4", "c5"), show='headings', style="mystyle.Treeview",yscrollcommand=trs.set)
    trs.config(command=tree.yview)

    tree.column("#1", anchor=tk.CENTER)

    tree.heading("#1", text="අයිතමයේ නම")

    tree.column("#2", anchor=tk.CENTER)

    tree.heading("#2", text="ප්‍රමාණය")

    tree.column("#3", anchor=tk.CENTER)

    tree.heading("#3", text="සාමන්‍ය මිල")

    tree.column("#4", anchor=tk.CENTER)

    tree.heading("#4", text="අපේ මිල")

    tree.column("#5", anchor=tk.CENTER)

    tree.heading("#5", text="එකතුව")

    tree.pack()

    #################remv#############################################
    def removeall():
        for recode in tree.get_children():
            tree.delete(recode)

    #########################################      save #############################
    def save():
        try:
            conn = sqlite3.connect('inv.db')
            sr = dis.get()
            print("HRI conn")
            c = conn.cursor()
            print("HRI cor")
            c.execute("DELETE FROM inv2 WHERE නම ='%s'" % sr)
            print("HRI cor")
            conn.commit()
            conn.close()

            def removeall():
                for recode in tree.get_children():
                    tree.delete(recode)

            nama = dis.get()
            q = float(qnt.get())
            np = float(normal.get())
            ap = float(apep.get())
            tot = float(total.get())

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("INSERT INTO inv2 VALUES ('%s','%s',%s,%s,%s)" % (nama, q, np, ap, tot))
            print('connected')
            conn.commit()
            conn.close()
            removeall()
            View()
            ' k'
            calc()
            normal.delete(0, END)
            apep.delete(0, END)
            dis.delete(0, END)
            qnt.delete(0, END)
            total.delete(0, END)

        except:
            def removeall():
                for recode in tree.get_children():
                    tree.delete(recode)

            nama = dis.get()
            q = float(qnt.get())
            np = float(normal.get())
            ap = float(apep.get())
            tot = float(total.get())

            conn = sqlite3.connect('inv.db')
            c = conn.cursor()
            print('connected')
            c.execute("INSERT INTO inv2 VALUES ('%s','%s',%s,%s,%s)" % (nama, q, np, ap, tot))
            print('connected')
            conn.commit()
            conn.close()
            removeall()
            View()
            calc()
            ' k'
            normal.delete(0, END)
            apep.delete(0, END)
            dis.delete(0, END)
            qnt.delete(0, END)
            total.delete(0, END)

    def select_2(event):
        try:
            dis.delete(0, END)
            selected = tree.focus()
            val = tree.item(selected, "values")
            dis.insert(0, val[0])

            qnt.delete(0, END)
            selected = tree.focus()
            val = tree.item(selected, "values")
            qnt.insert(0, val[1])

            normal.delete(0, END)
            selected = tree.focus()
            val = tree.item(selected, "values")
            normal.insert(0, val[2])

            apep.delete(0, END)
            selected = tree.focus()
            val = tree.item(selected, "values")
            apep.insert(0, val[3])

            total.delete(0, END)
            selected = tree.focus()
            val = tree.item(selected, "values")
            x_1 = (float(val[1]) * float(val[3]))
            total.insert(0, x_1)

        except:
            print("අහක ක්ලික් කරාම")
            normal.delete(0, END)
            apep.delete(0, END)
            dis.delete(0, END)
            qnt.delete(0, END)
            total.delete(0, END)










    def remover():

        conn = sqlite3.connect('inv.db')
        sr = dis.get()
        print("HRI conn")
        c = conn.cursor()
        print("HRI cor")
        c.execute("DELETE FROM inv3 WHERE name ='%s'" % sr)
        print("HRI cor")
        conn.commit()
        conn.close()

        x = tree.selection()[0]
        tree.delete(x)
        conn = sqlite3.connect('inv.db')
        sr = dis.get()
        print("HRI conn")
        c = conn.cursor()
        print("HRI cor")
        c.execute("DELETE FROM inv2 WHERE නම ='%s'" % sr)
        print("HRI cor")
        conn.commit()
        conn.close()
        removeall()
        View()
        bar.delete(0, END)
        dis.delete(0, END)
        qnt.delete(0, END)
        normal.delete(0, END)
        apep.delete(0, END)
        total.delete(0, END)
        list.destroy()
        inv()

    def select(event):

        conn = sqlite3.connect('item.db')
        sr = float(bar.get())

        print("HRI conn")
        c = conn.cursor()
        print("HRI cor")
        c.execute("SELECT NAME FROM item2 WHERE NUMBER ='%s'" % sr)
        rec = c.fetchone()[0]
        dis.insert(0, rec)

        print(rec)

        ###########
        c.execute("SELECT RPRICE FROM item2 WHERE NUMBER ='%s'" % sr)
        RET = float(c.fetchone()[0])

        normal.insert(0, RET)
        qnt.insert(0, 1)
        c.execute("SELECT DISCOUNT FROM item2 WHERE NUMBER ='%s'" % sr)
        DIS = float(c.fetchone()[0])
        print(DIS)
        print(RET)
        TOTL = RET - DIS
        apep.insert(0, TOTL)

        total.insert(0, TOTL)
        save()
        bar.delete(0, END)
        total.delete(0, END)

    def conf(event):
        try:
            q = float(qnt.get())
            nr = float(apep.get())
            TOTL = (q * nr)
            total.delete(0, END)
            total.insert(0, TOTL)
        except:
            print((""))

    View()

    def printing():
        try:
            document = docx.Document()



            import sqlite3

            connection = sqlite3.connect('inv.db')
            cursor = connection.execute('SELECT COUNT(නම) FROM inv2')
            pagecount = int(cursor.fetchone()[0])
            cursor = connection.execute('SELECT SUM(qt) FROM inv3')
            np = cursor.fetchone()[0]

            cursor = connection.execute('SELECT SUM(qr) FROM inv3')
            ap = cursor.fetchone()[0]
            print(pagecount)
            print(np)
            print(ap)
            labe = float(ap - np)

            print(labe)
            connection.commit()
            connection.close()

            y = int(pagecount) * 10 + 97
            print(y)

            today = date.today()

            t = time.localtime()
            current_time = time.strftime("%H:%M:%S", t)
            current_time = str(current_time)
            print(type(current_time))
            d2 = today.strftime("%B %d, %Y")

            ###########################################
            section = document.sections[0]
            section.page_height = Mm(y)
            section.page_width = Mm(80)
            section.left_margin = Mm(5)
            section.right_margin = Mm(1)
            section.top_margin = Mm(0)
            section.bottom_margin = Mm(0)
            section.header_distance = Mm(0)
            section.footer_distance = Mm(0)
            #############################################

            # Setting the Normal font works:
            paragraph = document.add_paragraph(str('           මෙත්මල් සුපර්'))

            font = paragraph.style.font
            font.size = Pt(20)
            font.bold = 1

            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            #########################################
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('head', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(10)
            obj_font.name = 'Iskoola Pota'

            ################################################

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(10)
            obj_font.name = 'Arial'

            ##################################################

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('PODI AKURU', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(8)
            obj_font.name = 'Arial'

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('big', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(14)
            obj_font.name = 'Times New Roman'

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('big1', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(12)
            obj_font.name = 'Times New Roman'

            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('small', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(6)
            obj_font.name = 'Arial'

            paragraph.add_run(
                "\n" + "              මීකණුව පාර හග්ගල්ල එල්ලක්කල" + "\n" + "                            0775143700",
                style='CommentsStyle')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)
            ##########################
            paragraph.add_run("\n" + "-------------------------------------------" + "\n", style='big')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)
            ################################################################################################
            paragraph.add_run(d2 + "        " + "*** WELCOME ***" + "         " + current_time, style='PODI AKURU')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)
            ################################################

            paragraph.add_run("\n" + "-------------------------------------------", style='big')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            ###############################################

            paragraph.add_run(
                "\n" + "නම/ප්‍රමාණය    සාමන්‍ය මිල     අප මිල    එකතුව" + "\n" + "-------------------------------------------------------------" + "\n",
                style='CommentsStyle')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            #####################################
            import sqlite3

            connection = sqlite3.connect('inv.db')
            cursor = connection.execute('select * from inv2')
            cursor = cursor.fetchall()
            print(cursor)

            for name in cursor:
                global baseekata_danna_Q_WPR
                global Total_QWpr
                global MULUWIKUNUM
                global shudda_labaya

                x = str(name[0])
                connection = sqlite3.connect('item.db')
                cursor = connection.execute("SELECT QUNTY from item2 WHERE NAME='%s'" % x)
                wcursor = connection.execute("SELECT QUNTY from item2 WHERE NAME='%s'" % x)
                it = float(cursor.fetchone()[0])
                print(it)
                print(x)
                q = float(name[1])
                g = float(name[2])
                print(g)
                z = float(name[3])
                print(g)
                xy = float(name[4])
                print(g)
                b = float(q)
                print(q)

                connection.commit()
                cursor.close()
                itb = it - b
                print(itb)
                ################################################################
                if q % 1 == 0:
                    q = int(q)
                    b = q
                b = q
                b = float(b)
                if b % 1 == 0:
                    b = int(b)
                b = len(str(b))
                ###################laba alaba ekata####################################
                wcursor = connection.execute("SELECT WPRICE from item2 WHERE NAME='%s'" % x)
                WPR = float(wcursor.fetchone()[0])
                print("wikunum mila  " + str(WPR))
                print("pramne" + str(q))
                Q_WPR = q * WPR
                baseekata_danna_Q_WPR = baseekata_danna_Q_WPR + Q_WPR
                Total_QWpr = Total_QWpr + Q_WPR

                print("අන්ක ගාන" + str(b) + x)

                if b == 1:
                    space1 = "                 "

                elif b == 2:

                    space1 = "                 "
                elif b == 3:
                    space1 = "                "

                elif b == 4:
                    space1 = "                 "

                elif b == 5:
                    space1 = "               "

                else:
                    space1 = "             "
                ##############################################################################################
                bx = g
                bx = int(bx)
                bx = len(str(bx))

                if bx == 1:

                    space11 = "                "
                elif bx == 2:
                    space11 = "              "

                elif bx == 3:
                    space11 = "            "

                elif bx == 4:
                    space11 = "          "

                elif bx == 5:
                    space11 = "        "

                else:
                    space11 = "      "
                print((bx))

                ################################################################## lada laba
                bxz = xy
                bxz = int(bx)
                bxz = len(str(bxz))

                if bxz == 1:

                    space1111 = "        "
                elif bxz == 2:
                    space1111 = "      "

                elif bxz == 3:
                    space1111 = "    "

                elif bxz == 4:
                    space1111 = "  "

                else:
                    space1111 = ""

                print((bxz))
                ##################################################################################

                ###################################################################

                #####################################################

                ##########################################################

                q = str(q)
                g = str(g)
                z = str(z)
                xy = str(xy)

                connection = sqlite3.connect('item.db')
                cursor = connection.cursor()
                sql_update_query = """Update item2 set QUNTY = ? where NAME = ?"""
                data = (itb, x)
                cursor.execute(sql_update_query, data)
                connection.commit()
                print("Record Updated successfully")
                cursor.close()
                ############################################################################space

                #################################################################################

                paragraph.add_run(
                    x + '\n' + q + space1 + g + space11 + z + space1111 + xy + '\n',
                    style='CommentsStyle')
                paragraph_format = paragraph.paragraph_format

            paragraph.add_run("\n" + "-------------------------------------------" + "\n", style='big')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            print("mulu okkom wikunum PRIWARY mila*q = dalalabe" + str(Total_QWpr))
            connection = sqlite3.connect('inv.db')
            abcursor = connection.execute('SELECT SUM(qt) FROM inv3')
            MULUWIKUNUM = abcursor.fetchone()[0]
            shudda_labaya = MULUWIKUNUM - Total_QWpr
            print(shudda_labaya)
            print("meka" + str(baseekata_danna_Q_WPR))

            #######################################space###################

            m = ap
            m = int(m)
            m = len(str(m))

            if m == 1:
                space1 = "               "

            elif m == 2:

                space1 = "             "
            elif m == 3:
                space1 = "           "

            elif m == 4:
                space1 = "         "

            elif m == 5:
                space1 = "       "

            else:
                space1 = "     "

            mx = np
            mx = int(mx)
            mx = len(str(mx))

            if mx == 1:

                space11 = "                       "
            elif mx == 2:
                space11 = "                     "

            elif mx == 3:
                space11 = "                   "

            elif mx == 4:
                space11 = "                 "

            elif mx == 5:
                space11 = "               "

            else:
                space11 = "             "
            print((mx))

            ################################################################## lada laba
            mxz = labe
            mxz = int(mxz)
            mxz = len(str(mxz))

            if mxz == 1:

                space1111 = "                      "
            elif mxz == 2:
                space1111 = "                    "

            elif mxz == 3:
                space1111 = "                  "

            elif mxz == 4:
                space1111 = "                "

            elif mxz == 5:
                space1111 = "              "

            else:
                space1111 = "            "
            print((mxz))

            ##########################################

            #####################################space######################

            paragraph.add_run("                         මුලු වටිනාකම" + space1 + str(
                ap) + "\n" + "                         අපේ මිල" + space11 + str(
                np), style='big1')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            paragraph.add_run("\n" + "-------------------------------------------" + "\n", style='big')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            paragraph.add_run("    *** සුබ දවසක් නැවත එන්න ***" + "\n", style='big')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            paragraph.add_run("\n" + "                        --System Created By KeZaRa 0768000534--" + "\n",
                              style="small")
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(1)
            paragraph.line_spacing = Mm(0.2)

            messagebox.showinfo(" රැදී සිටින්න", "ප්‍රින්ට් වෙමින් පවතී.(තප්පර 10)")

            document.save('test.docx')

            # Converting docx present in the same folder
            # as the python file
            convert("test.docx")

            # Converting docx specifying both the input
            # and output paths
            convert("test.docx", "Mine.pdf")

            # Notice that the output filename need not be
            # the same as the docx

            # Bulk Conversion
            convert("test.docx")
            # fd = os.startfile("test.pdf", "print")
            current_time = str(current_time.replace(":", '-', ))
            d2 = str(d2)

            # os.rename("1.pdf",str(d2)+"___"+str(current_time)+".pdf")
            original = r'E:\nethmal\test.pdf'
            target = r'E:\nethmal\pdf\%s' % ((d2) + "___" + (current_time) + ".pdf")

            shutil.copyfile(original, target)

            removeall()

            conn = sqlite3.connect('inv.db')
            sr = dis.get()
            print("HRI conn")
            c = conn.cursor()
            print("HRI cor")
            c.execute("DELETE FROM inv3")
            c.execute("DELETE FROM inv2")
            print("HRI cor")
            conn.commit()
            conn.close()
            dine_iwarai()
            list.destroy()
            inv()
        except:
            messagebox.showerror("කණගාටුයි", "ගනුදෙනුව සිදු කල නොහැක නැවත උත්සහ කරන්න")

    def dine_iwarai():
        today = date.today()
        conn = sqlite3.connect('inv.db')
        c = conn.cursor()
        c.execute("INSERT INTO labaya VALUES ('%s','%s','%s')" % (today, MULUWIKUNUM, shudda_labaya))
        conn.commit()
        conn.close()



    def dine_iwarai2():

        conn = sqlite3.connect('inv.db')
        c = conn.cursor()
        print('connected')
        c.execute("SELECT SUM(muluwikunum) FROM labaya ")
        x = float(c.fetchone()[0])

        conn = sqlite3.connect('inv.db')
        c = conn.cursor()
        print('connected')
        c.execute("SELECT SUM(shuddalabaya) FROM labaya ")
        yy = float(c.fetchone()[0])

        conn = sqlite3.connect('inv.db')





        today = date.today()
        conn = sqlite3.connect('inv.db')
        c = conn.cursor()
        c.execute("INSERT INTO labaya2 VALUES ('%s','%s','%s')" % (today, x, yy))
        conn.commit()
        conn.close()
        list.destroy()
        c.execute("DELETE FROM labaya")
        print("HRI cor")
        conn.commit()
        conn.close()


    ok = PhotoImage(file="img/ok.png")
    b1 = Button(list, image=ok, highlightthickness=0, bd=0, compound=LEFT, command=save)
    b1.place(x=978, y=65)

    cansel = PhotoImage(file="img/cansel.png")
    b2 = Button(list, image=cansel, highlightthickness=0, bd=0, compound=LEFT, command=remover)
    b2.place(x=978, y=137)

    oka = PhotoImage(file="img/printing.png")
    b1 = Button(list, image=oka, highlightthickness=0, bd=0, compound=LEFT, command=printing)
    b1.place(x=980, y=192)

    sav = PhotoImage(file="img/ස.png")
    b3 = Button(list, image=sav, highlightthickness=0, bd=0, compound=LEFT, command=dine_iwarai2)
    b3.place(x=403, y=621)

    tree.bind("<ButtonRelease-1>", select_2)
    bar.bind('<Return>', select)
    qnt.bind('<KeyRelease>', conf)
    bar.delete(0, END)

    list.mainloop()



######################################profit######################

bgr = PhotoImage(file="img/laba.png")
def profit():
    pr = Toplevel()
    pr.title("NETHMAL SUPERT MARKET BILLING SYSTEM")

    pr.iconbitmap("img/a.ico")
    pr.config(background='black')
    pr.geometry('1024x768')
    pr.resizable(width=False, height=False)
    bgr = PhotoImage(file="img/laba.png")
    l1 = Label(pr, image=bgr)
    l1.place(x=0, y=0)

    style = ttk.Style()
    style.configure("mystyle.Treeview", highlightthickness=0, bd=0,
                    font=('Calibri', 11))  # Modify the font of the body
    style.configure("mystyle.Treeview", background='black')
    style.configure("mystyle.Treeview.Heading", font=('Calibri', 13, 'bold'))  # Modify the font of the headings
    style.layout("mystyle.Treeview", [('mystyle.Treeview.treearea', {'sticky': 'nswe'})])  # Remove the borders

    fram = Frame(pr)
    fram.place(x=60, y=10)
    trs = Scrollbar(fram)
    trs.pack(side=LEFT, fill=Y)
    tree = ttk.Treeview(fram, column=("c1", "c2", "c3"), show='headings', style="mystyle.Treeview",
                        yscrollcommand=trs.set)
    trs.config(command=tree.yview)

    tree.column("#1", anchor=tk.CENTER)
    tree.column("#1", minwidth=0, width=500, stretch=NO)

    tree.heading("#1", text="දිනය")

    tree.column("#2", anchor=tk.CENTER)

    tree.heading("#2", text="මුලු විකුනුම්")

    tree.column("#3", anchor=tk.CENTER)

    tree.heading("#3", text="ශුද්ධ ලාබය")

    style = ttk.Style(pr)
    # set ttk theme to "clam" which support the fieldbackground option
    style.theme_use("clam")
    style.configure("Treeview", background="black",
                    fieldbackground="black", foreground="white")

    con1 = sqlite3.connect("inv.db")

    cur1 = con1.cursor()

    cur1.execute("SELECT * FROM labaya")

    rows = cur1.fetchall()

    for row in rows:
        print(row)

        tree.insert("", tk.END, values=row)

    con1.close()
    ' k'

    tree.pack()

    ############################################################
    def bar():
        connection = sqlite3.connect('inv.db')
        cursor = connection.execute('select date from labaya2')
        date = ""

        date = cursor.fetchall()

        #####################################
        connection = sqlite3.connect('inv.db')
        cursor = connection.execute('select shuddawikunum from labaya2')
        muluwikunum = ""

        rec = cursor.fetchall()
        print(rec)
        li = []
        for name in rec:
            labe = str(name[0])
            muluwikunum = labe
            li.append(muluwikunum)

        for i in range(0, len(li)):
            li[i] = float(li[i])
        print(li)

        muluwikunum = muluwikunum + "methan"
        print(muluwikunum)
        muluwikunum = muluwikunum.replace(',methan', '')
        print(muluwikunum)

        print(li)

        langs = []
        # date=str(java,python)
        langs.insert(0, date)
        print(langs)

        from matplotlib import pyplot as plt
        import numpy as np
        fig = plt.figure()
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis('equal')
        langs = date

        print(name)

        students = li

        ax.pie(students, labels=langs, autopct='%1.2f%%')
        plt.show()

    def dine_ekka_labe():
        style = ttk.Style()
        style.configure("mystyle.Treeview", highlightthickness=0, bd=0,
                        font=('Calibri', 11))  # Modify the font of the body
        style.configure("mystyle.Treeview", background='black')
        style.configure("mystyle.Treeview.Heading", font=('Calibri', 13, 'bold'))  # Modify the font of the headings
        style.layout("mystyle.Treeview", [('mystyle.Treeview.treearea', {'sticky': 'nswe'})])  # Remove the borders

        fram = Frame(pr)
        fram.place(x=60, y=500)
        trs = Scrollbar(fram)
        trs.pack(side=LEFT, fill=Y)
        tree = ttk.Treeview(fram, column=("c1", "c2", "c3"), show='headings', style="mystyle.Treeview",
                            yscrollcommand=trs.set)
        trs.config(command=tree.yview)

        tree.column("#1", anchor=tk.CENTER)
        tree.column("#1", minwidth=0, width=500, stretch=NO)

        tree.heading("#1", text="දිනය")

        tree.column("#2", anchor=tk.CENTER)

        tree.heading("#2", text="මුලු විකුනුම්")

        tree.column("#3", anchor=tk.CENTER)

        tree.heading("#3", text="ශුද්ධ ලාබය")

        style = ttk.Style(pr)
        # set ttk theme to "clam" which support the fieldbackground option
        style.theme_use("clam")
        style.configure("Treeview", background="black",
                        fieldbackground="black", foreground="white")

        con1 = sqlite3.connect("inv.db")

        cur1 = con1.cursor()

        cur1.execute("SELECT * FROM labaya2")

        rows = cur1.fetchall()

        for row in rows:
            print(row)

            tree.insert("", tk.END, values=row)

        con1.close()
        ' k'

        tree.pack()

    dine_ekka_labe()


    sav1 = PhotoImage(file="img/BAR.png")
    b3 = Button(pr, image=sav1, highlightthickness=0, bd=0, compound=LEFT, command=bar)
    b3.place(x=503, y=293)

    pr.mainloop()


#################################################################################


#########ADD BUTTON
imgadd = PhotoImage(file="img/add.png")
b1 = Button(root, image=imgadd, highlightthickness=0, bd=0, compound=LEFT, command=entru)
b1.place(x=67, y=21, )

imginv = PhotoImage(file="img/invoice.png")
b2 = Button(root, image=imginv, highlightthickness=0, bd=0, compound=LEFT, command=inv)
b2.place(x=451, y=26, )

prof = PhotoImage(file="img/profbut.png")
b3 = Button(root, image=prof, highlightthickness=0, bd=0, compound=LEFT, command=profit)
b3.place(x= 66, y=410, )

root.mainloop()
